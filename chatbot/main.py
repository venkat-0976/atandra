"""
FastAPI backend for document chat with Gemini AI - Directory-based version

This script:
1. Scans a specified directory for various document types (PDF, CSV, Word, Markdown, JSON, etc.)
2. Processes and extracts text content from these documents
3. Uses Google's Gemini AI to answer questions based on document content
4. Provides a REST API for frontend integration
"""

from fastapi import FastAPI, HTTPException, Query, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Dict, Any, Set
import google.generativeai as genai
import os
import io
import json
import PyPDF2
import pandas as pd
import docx
import markdown
from dotenv import load_dotenv
from pathlib import Path
import uvicorn
import logging
import asyncio
import time

# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables from .env file
load_dotenv()

# Get API key from environment variable
api_key = os.getenv("GEMINI_API_KEY")

# Initialize FastAPI app
app = FastAPI(
    title="Directory-Based Document Chat using Gemini AI",
    description="Chat with documents in a specified directory using Gemini AI",
    version="1.0.0",
)

# Configure CORS to allow requests from the frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Data models
class DocumentInfo(BaseModel):
    """Information about a processed document"""
    path: str
    filename: str
    extension: str
    content_preview: str
    char_count: int
    last_modified: float

class ChatRequest(BaseModel):
    """Request for chatting with documents"""
    query: str  # Changed from 'question' to 'query' to match frontend
    max_tokens: int = 500

class ChatResponse(BaseModel):
    """Response from Gemini AI"""
    answer: str
    sources: List[Dict[str, Any]] = []  # Added default empty list and changed from referenced_docs to sources

class Message(BaseModel):
    """Chat message"""
    role: str
    content: str

# Global settings
DEFAULT_DOCS_DIR = "Documents"  # Default directory to scan for documents
SUPPORTED_EXTENSIONS = {
    ".pdf": "PDF Document",
    ".csv": "CSV Spreadsheet",
    ".docx": "Word Document",
    ".doc": "Word Document",
    ".txt": "Text File",
    ".md": "Markdown File",
    ".json": "JSON File",
    ".xml": "XML File",
    ".html": "HTML File",
    ".htm": "HTML File",
    ".xlsx": "Excel Spreadsheet",
    ".xls": "Excel Spreadsheet",
}

# In-memory document cache
document_cache = {}
last_scan_time = 0
SCAN_INTERVAL = 30  # Rescan directory every 60 seconds

# In-memory chat history
chat_history = []

# Function to extract text from a PDF file
def extract_text_from_pdf(file_path):
    """Extract text content from a PDF file"""
    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            total_pages = len(pdf_reader.pages)
            
            text = f"[PDF: {os.path.basename(file_path)}]\n\n"
            for page_num in range(total_pages):
                # Extract text from page
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():  # Only add non-empty pages
                    text += f"--- Page {page_num + 1} ---\n"
                    text += page_text + "\n\n"
            
            return text
    except Exception as e:
        logger.error(f"Error extracting PDF {file_path}: {str(e)}")
        return f"[Error extracting PDF {os.path.basename(file_path)}: {str(e)}]"

# Function to extract data from a CSV file
def extract_data_from_csv(file_path):
    """Extract data from a CSV file"""
    try:
        # Parse CSV with pandas
        df = pd.read_csv(file_path)
        
        # Convert to string representation
        csv_content = f"[CSV: {os.path.basename(file_path)}]\n"
        csv_content += f"Rows: {len(df)}, Columns: {', '.join(df.columns.tolist())}\n\n"
        csv_content += df.to_string(index=True)
        
        return csv_content
    except Exception as e:
        logger.error(f"Error processing CSV {file_path}: {str(e)}")
        return f"[Error processing CSV {os.path.basename(file_path)}: {str(e)}]"

# Function to extract text from a Word document
def extract_text_from_docx(file_path):
    """Extract text from a Word document"""
    try:
        doc = docx.Document(file_path)
        text = f"[Word Document: {os.path.basename(file_path)}]\n\n"
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                text += " | ".join(row_text) + "\n"
            text += "\n"
            
        return text
    except Exception as e:
        logger.error(f"Error extracting Word document {file_path}: {str(e)}")
        return f"[Error extracting Word document {os.path.basename(file_path)}: {str(e)}]"

# Function to extract text from a Markdown file
def extract_text_from_markdown(file_path):
    """Extract text from a Markdown file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to plain text (simplified)
        # You might want to use a more sophisticated approach for production
        text = f"[Markdown: {os.path.basename(file_path)}]\n\n"
        text += md_content  # Keep the markdown content as is
        
        return text
    except Exception as e:
        logger.error(f"Error processing Markdown {file_path}: {str(e)}")
        return f"[Error processing Markdown {os.path.basename(file_path)}: {str(e)}]"

# Function to extract data from a JSON file
def extract_data_from_json(file_path):
    """Extract data from a JSON file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        # Convert to string representation
        text = f"[JSON: {os.path.basename(file_path)}]\n\n"
        text += json.dumps(json_data, indent=2)
        
        return text
    except Exception as e:
        logger.error(f"Error processing JSON {file_path}: {str(e)}")
        return f"[Error processing JSON {os.path.basename(file_path)}: {str(e)}]"

# Function to extract text from a plain text file
def extract_text_from_txt(file_path):
    """Extract text from a plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        text = f"[Text File: {os.path.basename(file_path)}]\n\n"
        text += content
        
        return text
    except UnicodeDecodeError:
        try:
            # Try again with a different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                content = f.read()
            
            text = f"[Text File: {os.path.basename(file_path)}]\n\n"
            text += content
            
            return text
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            return f"[Error processing text file {os.path.basename(file_path)}: {str(e)}]"
    except Exception as e:
        logger.error(f"Error processing text file {file_path}: {str(e)}")
        return f"[Error processing text file {os.path.basename(file_path)}: {str(e)}]"

# Function to extract data from an Excel file
def extract_data_from_excel(file_path):
    """Extract data from an Excel file"""
    try:
        # Read all sheets
        df_dict = pd.read_excel(file_path, sheet_name=None)
        
        # Convert to string representation
        excel_content = f"[Excel: {os.path.basename(file_path)}]\n\n"
        
        for sheet_name, df in df_dict.items():
            excel_content += f"Sheet: {sheet_name}\n"
            excel_content += f"Rows: {len(df)}, Columns: {', '.join(df.columns.tolist())}\n"
            
            # Print first 100 rows max to avoid extremely large strings
            preview_df = df.head(100)
            excel_content += preview_df.to_string(index=True) + "\n\n"
            
            if len(df) > 100:
                excel_content += f"[...{len(df) - 100} more rows not shown...]\n\n"
        
        return excel_content
    except Exception as e:
        logger.error(f"Error processing Excel file {file_path}: {str(e)}")
        return f"[Error processing Excel file {os.path.basename(file_path)}: {str(e)}]"

# Extract content from an HTML file
def extract_text_from_html(file_path):
    """Extract text from an HTML file"""
    try:
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract text (remove script and style elements)
        for script in soup(["script", "style"]):
            script.extract()
        
        text = soup.get_text()
        
        # Format the text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return f"[HTML: {os.path.basename(file_path)}]\n\n" + text
    except ImportError:
        logger.warning("BeautifulSoup not installed. Using raw HTML.")
        return extract_text_from_txt(file_path)
    except Exception as e:
        logger.error(f"Error processing HTML file {file_path}: {str(e)}")
        return f"[Error processing HTML file {os.path.basename(file_path)}: {str(e)}]"

# Function to process a document and extract its content
def process_document(file_path):
    """Process a document and extract its content based on file extension"""
    _, extension = os.path.splitext(file_path)
    extension = extension.lower()
    
    # Process based on file extension
    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif extension == '.csv':
        return extract_data_from_csv(file_path)
    elif extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif extension == '.md':
        return extract_text_from_markdown(file_path)
    elif extension == '.json':
        return extract_data_from_json(file_path)
    elif extension in ['.txt', '.log', '.sql', '.py', '.js', '.css', '.xml']:
        return extract_text_from_txt(file_path)
    elif extension in ['.html', '.htm']:
        return extract_text_from_html(file_path)
    elif extension in ['.xlsx', '.xls']:
        return extract_data_from_excel(file_path)
    else:
        logger.warning(f"Unsupported file type: {extension} for {file_path}")
        return f"[Unsupported file type: {extension} for {os.path.basename(file_path)}]"

# Scan directory for documents
def scan_documents(directory_path=DEFAULT_DOCS_DIR):
    """Scan directory for documents and process them"""
    global document_cache, last_scan_time
    
    # Check if directory exists, if not create it
    if not os.path.exists(directory_path):
        os.makedirs(directory_path)
        logger.info(f"Created documents directory: {directory_path}")
    
    # Check if we should rescan based on time interval
    current_time = time.time()
    if current_time - last_scan_time < SCAN_INTERVAL and document_cache:
        return document_cache
    
    # Clear existing cache for clean scan
    document_cache = {}
    files_processed = 0
    
    logger.info(f"Scanning directory for documents: {directory_path}")
    
    for root, _, files in os.walk(directory_path):
        for file in files:
            file_path = os.path.join(root, file)
            _, extension = os.path.splitext(file_path)
            extension = extension.lower()
            
            # Check if file extension is supported
            if extension in SUPPORTED_EXTENSIONS:
                try:
                    # Get file metadata
                    file_stats = os.stat(file_path)
                    last_modified = file_stats.st_mtime
                    
                    # Only reprocess if file is new or modified
                    if file_path in document_cache and document_cache[file_path]['last_modified'] >= last_modified:
                        files_processed += 1
                        continue
                    
                    # Process the document
                    content = process_document(file_path)
                    
                    # Add to cache
                    document_cache[file_path] = {
                        'path': file_path,
                        'filename': file,
                        'extension': extension,
                        'content': content,
                        'content_preview': content[:200] + "..." if len(content) > 200 else content,
                        'char_count': len(content),
                        'last_modified': last_modified
                    }
                    
                    files_processed += 1
                    logger.info(f"Processed: {file} ({extension})")
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {str(e)}")
    
    last_scan_time = current_time
    logger.info(f"Document scanning complete. Found {files_processed} files.")
    return document_cache

# Function to interact with Gemini API
async def chat_with_gemini(prompt, document_content, api_key):
    """Interact with Gemini API to get answers based on document content"""
    if not api_key:
        raise HTTPException(status_code=500, detail="Error: No Gemini API key found in environment variables")
    
    try:
        # Configure the API
        genai.configure(api_key=api_key)
        
        # Set up the model
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Create a system prompt - MODIFIED to remove reference to document sources
        system_prompt = """
        You are a helpful assistant for Atandra Energy that answers questions based on the provided document content.
        Analyze the document content carefully and provide accurate responses.
        If the answer cannot be found in the documents, clearly state that.
        Format your responses in a clear, concise manner.
        DO NOT mention which documents you referenced in your response.
        DO NOT include any section about "Documents Referenced" or similar information at the end of your response.
        """
        
        # Generate response considering the document context - Modified to remove instructions about document references
        full_prompt = f"{system_prompt}\n\nDOCUMENT CONTENT:\n{document_content}\n\nUSER QUESTION: {prompt}\n\nPlease provide a helpful response based on the document content."
        
        response = model.generate_content(full_prompt)
        return response.text
    except Exception as e:
        logger.error(f"Error from Gemini AI: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error from Gemini AI: {str(e)}")

@app.get("/")
async def root():
    """Root endpoint to check if the API is running"""
    return {
        "status": "online", 
        "message": "Directory-Based Document Chat using Gemini AI",
        "docs_directory": DEFAULT_DOCS_DIR,
        "supported_formats": list(SUPPORTED_EXTENSIONS.keys())
    }

@app.get("/api/documents")
async def list_documents(directory: Optional[str] = Query(None)):
    """List all documents in the specified directory"""
    docs_dir = directory or DEFAULT_DOCS_DIR
    documents = scan_documents(docs_dir)
    
    doc_list = []
    for doc_path, doc_info in documents.items():
        doc_list.append({
            "path": doc_info['path'],
            "filename": doc_info['filename'],
            "extension": doc_info['extension'],
            "content_preview": doc_info['content_preview'],
            "char_count": doc_info['char_count'],
            "last_modified": doc_info['last_modified']
        })
    
    return {"documents": doc_list, "count": len(doc_list)}

@app.post("/api/scan")
async def force_scan(directory: Optional[str] = Query(None)):
    """Force a rescan of the documents directory"""
    global last_scan_time
    last_scan_time = 0  # Reset scan time to force rescan
    docs_dir = directory or DEFAULT_DOCS_DIR
    documents = scan_documents(docs_dir)
    
    return {
        "message": "Document scan complete",
        "documents_found": len(documents),
        "directory": docs_dir
    }

@app.get("/api/document/{filename}")
async def get_document_content(filename: str, directory: Optional[str] = Query(None)):
    """Get the content of a specific document by filename"""
    docs_dir = directory or DEFAULT_DOCS_DIR
    documents = scan_documents(docs_dir)
    
    # Find the document by filename
    for doc_path, doc_info in documents.items():
        if doc_info['filename'] == filename:
            return {
                "filename": doc_info['filename'],
                "content": doc_info['content'],
                "char_count": doc_info['char_count']
            }
    
    raise HTTPException(status_code=404, detail=f"Document not found: {filename}")

@app.post("/api/chat", response_model=ChatResponse)
async def api_chat(request: ChatRequest, directory: Optional[str] = Query(None)):
    """Chat with documents using Gemini AI (API endpoint)"""
    return await process_chat_request(request, directory)

# Add a new endpoint at /chat to match the frontend request URL
@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, directory: Optional[str] = Query(None)):
    """Chat with documents using Gemini AI (direct endpoint for frontend compatibility)"""
    return await process_chat_request(request, directory)

# Shared function to process chat requests from either endpoint
async def process_chat_request(request: ChatRequest, directory: Optional[str] = None):
    """Process a chat request from any endpoint"""
    docs_dir = directory or DEFAULT_DOCS_DIR
    documents = scan_documents(docs_dir)
    
    if not documents:
        raise HTTPException(status_code=400, detail=f"No documents found in directory: {docs_dir}")
    
    if not api_key:
        raise HTTPException(status_code=500, detail="No Gemini API key found. Please set GEMINI_API_KEY in the .env file")
    
    # Combine all document content
    all_docs_content = ""
    for doc_path, doc_info in documents.items():
        # Add document separator for clarity
        all_docs_content += f"\n\n{'=' * 30}\n"
        all_docs_content += f"DOCUMENT: {doc_info['filename']}\n"
        all_docs_content += f"{'=' * 30}\n\n"
        all_docs_content += doc_info['content']
    
    # Get AI response
    ai_response = await chat_with_gemini(request.query, all_docs_content, api_key)
    
    # Add to chat history
    chat_history.append({"role": "user", "content": request.query})
    chat_history.append({"role": "assistant", "content": ai_response})
    
    # We'll keep the sources list empty since we don't want to show document references
    return ChatResponse(answer=ai_response, sources=[])

@app.get("/api/history")
async def get_chat_history():
    """Get the chat history"""
    return {"messages": chat_history}

@app.delete("/api/history")
async def clear_chat_history():
    """Clear the chat history"""
    global chat_history
    chat_history = []
    return {"message": "Chat history cleared"}

# Run the server
if __name__ == "__main__":    # Initial document scan at startup
    scan_documents()
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)